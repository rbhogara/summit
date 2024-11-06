import streamlit as st
import io
import docx
from langchain_community.vectorstores import Chroma
from langchain_community import embeddings
from langchain_community.llms import Ollama
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from collections import namedtuple


#Extracts text from a DOCX file
def extract_text_from_docx(file_path):
    try:
        with open(file_path, "rb") as f:
            doc = docx.Document(f)
            text = ""
            for para in doc.paragraphs:
                text += para.text
        return text
    except Exception as e:
        st.error(str(e))
        return None



#Processes the uploaded files to extract text and generate embeddings
def process_input(file_paths):
    try:
        docs_list = []
        for file_path in file_paths:
            if file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                st.error(f"Unsupported file type: {file_path}")
                continue
            if text:
                docs_list.append(text)
        #split data into chunks
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(" ".join(docs_list))

        #create Documents from chunks
        Document = namedtuple('Document', ['page_content', 'metadata'])
        documents = [Document(page_content=chunk, metadata={}) for chunk in chunks]

        #Generate embeddings with nomic-embed-text model and store them in a vectorstore
        vectorstore = Chroma.from_documents(
            documents=documents,
            collection_name="rag-chroma",
            embedding=embeddings.OllamaEmbeddings(model='nomic-embed-text'),
        )

        return vectorstore.as_retriever()
    except Exception as e:
        st.error("Error processing input files.")
        st.error(str(e))
        return None

#Answers a question based on the context retrieved from the documents
def answer_question(question, retriever):
    try:
        model_local = Ollama(model="llama3")
        after_rag_template = """Answer the question based only on the following context:
        {context}
        Question: {question}
        """
        after_rag_prompt = ChatPromptTemplate.from_template(after_rag_template)
        after_rag_chain = (
            {"context": retriever, "question": RunnablePassthrough()}
            | after_rag_prompt
            | model_local
            | StrOutputParser()
        )
        return after_rag_chain.invoke(question)
    except Exception as e:
        st.error("Error processing question.")
        st.error(str(e))
        return None

#Main function to run the Streamlit app
def main():
    st.title("Security Summit")
    file_paths = [
        "recipe.docx"
    ]
    
    retriever = process_input(file_paths)
    
    with st.spinner('Generating Embeddings...'):
        if retriever:
            st.success('Embeddings generated successfully!')
            

    question_input = st.text_input("Question")

    if st.button('Query Documents'):
        with st.spinner('Processing...'):
            if retriever:
                answer = answer_question(question_input, retriever)
                st.text_area("Answer", value=answer, height=300, disabled=True)
            else:
                st.error("Failed to query documents.")

if __name__ == "__main__":
    main()
